import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Paths to actual images in the workspace
IMAGE_PATHS = {
    'login': r'e:\eventregis\docs\images\actual_login.png',
    'dashboard': r'e:\eventregis\docs\images\actual_dashboard.png',
    'view': r'e:\eventregis\docs\images\actual_registration_view.png'
}

# Custom styling helpers
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, space_before=Pt(12), space_after=Pt(6)):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = space_before
    heading.paragraph_format.space_after = space_after
    heading.paragraph_format.keep_with_next = True
    
    # Apply Mường Thanh Premium Burgundy color to headings
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(139, 0, 0)  # Burgundy
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(163, 0, 0)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        
    return heading

def add_paragraph_styled(doc, text="", style='Normal', space_after=Pt(6), line_spacing=1.15):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
    return p

def add_bullet_item(doc, bold_prefix, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.font.name = 'Arial'
        run_bold.font.size = Pt(10.5)
        run_bold.bold = True
    
    run_text = p.add_run(text)
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(10.5)
    return p

def add_step_paragraph(doc, step_name, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25)
    
    run_step = p.add_run(f"{step_name}: ")
    run_step.font.name = 'Arial'
    run_step.font.size = Pt(10.5)
    run_step.bold = True
    
    run_text = p.add_run(text)
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(10.5)
    return p

def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    # Shading background XML
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading)
    
    # Left border XML
    pBrd = OxmlElement('w:pBrd')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24')  # 3pt width
    left_border.set(qn('w:space'), '8')
    left_border.set(qn('w:color'), '8B0000')  # Burgundy border
    pBrd.append(left_border)
    p._p.get_or_add_pPr().append(pBrd)
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.italic = True
    return p

def add_centered_picture(doc, path, caption, width_in=5.8):
    if not path or not os.path.exists(path):
        add_paragraph_styled(doc, f"[Ảnh minh họa không khả dụng: {path}]", style='Normal')
        return
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    
    p_img.add_run().add_picture(path, width=Inches(width_in))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(f"Hình minh họa: {caption}")
    run_cap.font.name = 'Arial'
    run_cap.font.size = Pt(9)
    run_cap.italic = True
    run_cap.font.color.rgb = RGBColor(128, 128, 128)

def format_grid_table(table, col_widths, headers, data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header formatting
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], '8B0000')  # Burgundy Header
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data row formatting
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = 'F9F9F9' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            
    # Set cell widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)

# Main document generation
def generate_docx():
    doc = Document()
    
    # Configure margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("TÀI LIỆU HƯỚNG DẪN SỬ DỤNG")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(139, 0, 0)
    run_title.bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(48)
    run_sub = p_sub.add_run("QUY TRÌNH ĐĂNG NHẬP, ĐĂNG KÝ SỰ KIỆN & PHÊ DUYỆT\n(EVENTREGIS)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    add_paragraph_styled(doc, "Phiên bản: 2.0 (Cập nhật 2026)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_styled(doc, "Tech Stack: Yii1 PHP Framework | Nền tảng Portal SSO").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== SECTION 4 ====================
    add_heading_styled(doc, "4. Đăng nhập hệ thống", level=1)
    
    add_heading_styled(doc, "4.1 Đăng nhập hệ thống qua Portal SSO", level=2)
    add_step_paragraph(doc, "Bước 1", "Mở trình duyệt Chrome và truy cập địa chỉ hệ thống EventRegis: http://event.mt:8080/.")
    add_step_paragraph(doc, "Bước 2", "Hệ thống kiểm tra phiên đăng nhập. Nếu chưa đăng nhập, trình duyệt sẽ tự động chuyển hướng sang trang đăng nhập Portal Mường Thanh. Người dùng nhập các thông tin xác thực bắt buộc bao gồm:")
    add_bullet_item(doc, "", "Tên đăng nhập (Email hoặc Số điện thoại)", level=1)
    add_bullet_item(doc, "", "Mật khẩu", level=1)
    add_step_paragraph(doc, "Bước 3", "Nhấn nút “Đăng nhập”.")
    add_paragraph_styled(doc, "Kết quả: Hệ thống Portal xác thực thông tin tài khoản, tự động sinh mã Token JWT và điều hướng ngược trở lại EventRegis. Hệ thống khởi tạo Session và đưa bạn vào màn hình Dashboard chính (http://event.mt:8080/admin/default/index).")
    add_centered_picture(doc, IMAGE_PATHS['login'], "Giao diện đăng nhập tích hợp Portal SSO chạy tại http://event.mt:8080/", width_in=5.8)
    
    add_heading_styled(doc, "4.2 Đăng nhập nhanh từ trang chủ Portal (Portal Dashboard)", level=2)
    add_paragraph_styled(doc, "*Áp dụng trong trường hợp tài khoản đã được đăng nhập sẵn trên trình duyệt Chrome.*")
    add_step_paragraph(doc, "Bước 1", "Mở trình duyệt và truy cập trang chủ Portal: https://portal.muongthanh.vn.")
    add_step_paragraph(doc, "Bước 2", "Tại giao diện màn hình chính Portal (hiển thị danh mục các ứng dụng được phân quyền), người dùng tìm và click chọn biểu tượng ứng dụng “Đăng ký Sự kiện” (hoặc “Event Regis”).")
    add_step_paragraph(doc, "Bước 3", "Hệ thống Portal tự động sinh mã JWT Token và thực hiện chuyển hướng trình duyệt thẳng tới hệ thống EventRegis.")
    add_paragraph_styled(doc, "Kết quả: Bạn được đưa trực tiếp vào màn hình làm việc chính của EventRegis trong chưa đầy 1 giây mà không cần nhập lại bất kỳ thông tin nào.")
    
    add_heading_styled(doc, "4.3 Đăng nhập nhanh từ trang chủ EventRegis (Silent Login)", level=2)
    add_paragraph_styled(doc, "*Áp dụng trong trường hợp tài khoản đã được đăng nhập sẵn trên trình duyệt Chrome.*")
    add_step_paragraph(doc, "Bước 1", "Nhập trực tiếp địa chỉ hệ thống EventRegis trên thanh địa chỉ Chrome: http://event.mt:8080/.")
    add_step_paragraph(doc, "Bước 2", "Tại màn hình chào mừng của EventRegis, người dùng click vào nút màu tím “Đăng nhập với Portal”.")
    add_step_paragraph(doc, "Bước 3", "Hệ thống tự động chuyển hướng sang Portal để kiểm tra phiên. Do trình duyệt đã lưu session của bạn, Portal sẽ xác thực tự động ngay lập tức mà không hiển thị màn hình điền Tên đăng nhập & Mật khẩu.")
    add_paragraph_styled(doc, "Kết quả: Hệ thống điều hướng ngược lại EventRegis kèm mã sso_token, thiết lập session và cho phép bạn làm việc ngay lập tức.")
    
    doc.add_page_break()
    
    # ==================== SECTION 5 ====================
    add_heading_styled(doc, "5. Đăng ký tham gia sự kiện (Dành cho Đại diện Đơn vị)", level=1)
    
    add_heading_styled(doc, "5.1 Xem danh sách đợt đăng ký sự kiện", level=2)
    add_step_paragraph(doc, "Bước 1", "Đăng nhập vào hệ thống EventRegis với tư cách là Đại diện Đơn vị.")
    add_step_paragraph(doc, "Bước 2", "Trên thanh điều hướng bên trái (Sidebar), tìm và click chọn danh mục “Đăng ký Sự kiện” (hoặc truy cập trực tiếp http://event.mt:8080/admin/registrations/admin).")
    add_paragraph_styled(doc, "Kết quả: Hệ thống hiển thị bảng danh sách các đợt đăng ký sự kiện đang có trên hệ thống, bao gồm các thông tin: Tên đợt đăng ký, Thời gian mở đăng ký, Thời gian đóng đăng ký, Trạng thái hoạt động.")
    add_centered_picture(doc, IMAGE_PATHS['dashboard'], "Giao diện Danh sách phiếu đăng ký đại biểu đơn vị tại http://event.mt:8080/admin/registrations/admin", width_in=5.8)
    
    add_heading_styled(doc, "5.2 Khởi tạo Phiếu Đăng ký", level=2)
    add_step_paragraph(doc, "Bước 1", "Tại bảng danh sách đợt đăng ký sự kiện đang mở, người dùng click vào nút “Khởi tạo Phiếu Đăng ký” tương ứng với sự kiện muốn tham gia.")
    add_step_paragraph(doc, "Bước 2", "Hệ thống tự động tạo một phiếu đăng ký mới gắn liền với đơn vị thành viên của bạn.")
    add_paragraph_styled(doc, "Kết quả: Phiếu đăng ký được khởi tạo thành công với trạng thái mặc định ban đầu là “draft” (Bản nháp). Người dùng được chuyển đến giao diện chi tiết phiếu đăng ký để bắt đầu khai báo dữ liệu.")
    add_callout(doc, "Lưu ý nghiệp vụ quan trọng: Mỗi đơn vị chỉ được phép khởi tạo và sở hữu đúng một phiếu đăng ký hoạt động trong mỗi đợt đăng ký sự kiện để đảm bảo tính tập trung dữ liệu.")
    
    add_heading_styled(doc, "5.3 Nhập danh sách Đại biểu & Tải lên hồ sơ pháp lý bắt buộc", level=2)
    add_step_paragraph(doc, "Bước 1", "Tại giao diện chi tiết Phiếu Đăng ký, click chọn tab “Danh sách Đại biểu” và nhấn nút “Thêm đại biểu”.")
    add_step_paragraph(doc, "Bước 2", "Lựa chọn một trong hai phương thức khai báo nhân sự:")
    add_bullet_item(doc, "Phương thức 1: Đồng bộ từ hệ thống SMILE (Khuyên dùng): ", "Nhập Tên hoặc Mã nhân sự vào ô tìm kiếm. Click chọn nhân sự phù hợp từ danh sách gợi ý. Hệ thống sẽ tự động điền các thông tin: Họ tên, Chức vụ hiện tại, Mã phòng ban. Ràng buộc: Chỉ cho phép đăng ký những nhân sự có ngày gia nhập đơn vị trước ngày 01/06/2026.", level=1)
    add_bullet_item(doc, "Phương thức 2: Tự nhập thông tin: ", "Tích chọn ô “Tự nhập thông tin” để điền thủ công thông tin nhân viên chính thức nếu chưa được cập nhật dữ liệu kịp thời trên SMILE.", level=1)
    add_step_paragraph(doc, "Bước 3", "Tại phần hồ sơ tài liệu đính kèm của đại biểu, người dùng tải lên đầy đủ 4 tệp tin bắt buộc:")
    add_bullet_item(doc, "Ảnh mặt trước CCCD: ", "Định dạng ảnh JPG/PNG, dung lượng tối đa 5MB.", level=1)
    add_bullet_item(doc, "Ảnh mặt sau CCCD: ", "Định dạng ảnh JPG/PNG, dung lượng tối đa 5MB.", level=1)
    add_bullet_item(doc, "Ảnh chân dung in thẻ (Portrait): ", "Bắt buộc ảnh rõ nét, phông nền sáng và có kích thước chính xác 530x530 pixel. Hệ thống validate kích thước ở phía server.", level=1)
    add_bullet_item(doc, "Scan Hợp đồng lao động: ", "File scan định dạng PDF hoặc ảnh JPG rõ chữ ký/dấu đỏ thể hiện nhân sự chính thức, dung lượng tối đa 10MB.", level=1)
    add_step_paragraph(doc, "Bước 4", "Click nút “Lưu lại” để hoàn tất thêm mới đại biểu.")
    add_callout(doc, "Lưu ý quan trọng: Quy tắc validate ảnh chân dung 530x530px là bắt buộc để ảnh in thẻ đại biểu vật lý đạt tiêu chuẩn, rõ nét, không bị kéo dãn.")
    
    add_heading_styled(doc, "5.4 Thiết lập Liên quân theo từng Nội dung (Ghép đội thi đấu)", level=2)
    add_paragraph_styled(doc, "*Áp dụng khi các đơn vị thành viên quy mô nhỏ cần ghép nhân sự với nhau để thành lập đội thi đấu các nội dung tập thể.*")
    add_step_paragraph(doc, "Bước 1", "Tại giao diện chi tiết Phiếu Đăng ký, click chọn tab “Liên quân” và nhấn nút “Gửi yêu cầu liên quân”.")
    add_step_paragraph(doc, "Bước 2", "Khai báo biểu mẫu yêu cầu ghép đội bao gồm các trường thông tin: Đơn vị đối tác liên kết và Nội dung muốn liên quân (Ví dụ: Bóng đá nam, Kéo co...).")
    add_step_paragraph(doc, "Bước 3", "Click nút “Gửi yêu cầu”.")
    add_paragraph_styled(doc, "Kết quả: Yêu cầu liên quân được tạo ở trạng thái “pending” (Chờ duyệt). Khi đơn vị đối tác đăng nhập và chọn “Đồng ý”, liên quân sẽ chính thức được kích hoạt (active).")
    
    add_heading_styled(doc, "5.5 Đăng ký hoạt động chi tiết cho Đại biểu", level=2)
    add_step_paragraph(doc, "Bước 1", "Tại giao diện chi tiết Phiếu Đăng ký, click chọn tab “Đăng ký hoạt động”.")
    add_step_paragraph(doc, "Bước 2", "Thực hiện đăng ký cho từng nội dung sự kiện:")
    add_bullet_item(doc, "Môn thi đấu tập thể (Đăng ký theo số lượng): ", "Tích chọn tham gia và nhập số lượng đội thi đấu. Không cần chọn chi tiết danh sách thành viên ở bước này.", level=1)
    add_bullet_item(doc, "Môn thi cá nhân/ nghiệp vụ (Đăng ký theo danh sách): ", "Click chọn nội dung thi đấu (Miss, văn nghệ, nghiệp vụ), click chọn đại biểu cụ thể trong danh sách đại biểu để gán vào nội dung thi đấu.", level=1)
    add_step_paragraph(doc, "Bước 3", "Click “Lưu thông tin đăng ký” để cập nhật hệ thống.")
    add_callout(doc, "Hệ thống tự động kiểm duyệt các ràng buộc: Mỗi đại biểu được đăng ký tham gia tối đa N môn thể thao root (max_sports_per_attendee, mặc định là 3). Đồng thời, đại biểu thi nghiệp vụ bắt buộc phải thuộc mã phòng ban SMILE hợp lệ nằm trong cấu hình competition_departments.")
    
    add_heading_styled(doc, "5.6 Nộp hồ sơ đăng ký cho HO", level=2)
    add_step_paragraph(doc, "Bước 1", "Rà soát kỹ lưỡng toàn bộ danh sách đại biểu, tệp hồ sơ đính kèm, thông tin liên quân và danh sách môn thi đã đăng ký.")
    add_step_paragraph(doc, "Bước 2", "Cuộn lên đầu trang chi tiết Phiếu Đăng ký và click nút “Nộp đăng ký”.")
    add_step_paragraph(doc, "Bước 3", "Xác nhận nộp trong hộp thoại cảnh báo hiện ra.")
    add_paragraph_styled(doc, "Kết quả: Trạng thái phiếu đăng ký chuyển từ “draft” sang “submitted” (Đã nộp). Hệ thống tự động kích hoạt chế độ Khóa chỉnh sửa (Read-only) của đơn vị thành viên để đảm bảo tính toàn vẹn dữ liệu trong suốt quá trình kiểm duyệt.")
    
    doc.add_page_break()
    
    # ==================== SECTION 6 ====================
    add_heading_styled(doc, "6. Quy Trình Kiểm Duyệt & Phê Duyệt (Dành cho Nhân sự HO)", level=1)
    
    add_heading_styled(doc, "6.1 Tiếp nhận & Thẩm định Hồ sơ Đại biểu", level=2)
    add_step_paragraph(doc, "Bước 1", "Đăng nhập hệ thống quản trị EventRegis bằng tài khoản Nhân sự HO (HR HO).")
    add_step_paragraph(doc, "Bước 2", "Truy cập mục “Kiểm duyệt Đăng ký” (hoặc http://event.mt:8080/admin/registrations/admin). Tìm và mở chi tiết phiếu đăng ký của đơn vị đang có trạng thái “submitted”.")
    add_step_paragraph(doc, "Bước 3", "Cuộn xuống danh sách đại biểu, click xem chi tiết từng nhân sự để kiểm duyệt trực quan 4 tệp hồ sơ pháp lý bắt buộc: ảnh chân dung 530x530px, ảnh 2 mặt CCCD và tệp scan HĐLĐ để xác minh đại biểu.")
    add_centered_picture(doc, IMAGE_PATHS['view'], "Giao diện Chi tiết Phiếu đăng ký của HR HO tại http://event.mt:8080/admin/registrations/view/id/2", width_in=5.8)
    
    add_heading_styled(doc, "6.2 Từ chối Phiếu Đăng ký (Reject)", level=2)
    add_paragraph_styled(doc, "*Áp dụng khi phát hiện hồ sơ của ít nhất một đại biểu trong danh sách bị lỗi hoặc không đạt tiêu chuẩn nghiệp vụ.*")
    add_step_paragraph(doc, "Bước 1", "Tại giao diện chi tiết phiếu đăng ký cần từ chối, HR HO click nút “Từ chối”.")
    add_step_paragraph(doc, "Bước 2", "Trong hộp thoại hiện ra, bắt buộc phải nhập lý do từ chối chi tiết (Ví dụ: “Ảnh chân dung đại biểu Nguyễn Văn A bị mờ”).")
    add_step_paragraph(doc, "Bước 3", "Click nút “Xác nhận từ chối”.")
    add_paragraph_styled(doc, "Kết quả: Trạng thái phiếu đăng ký chuyển sang “rejected” (Từ chối). Hệ thống tự động mở khóa quyền chỉnh sửa trên phiếu đăng ký để đơn vị cập nhật hồ sơ lỗi và nộp lại.")
    
    add_heading_styled(doc, "6.3 Phê duyệt Phiếu Đăng ký (Approve)", level=2)
    add_paragraph_styled(doc, "*Áp dụng khi toàn bộ danh sách đại biểu và hồ sơ đính kèm của đơn vị đã đạt yêu cầu nghiệp vụ.*")
    add_step_paragraph(doc, "Bước 1", "Tại giao diện chi tiết phiếu đăng ký của đơn vị, HR HO click nút “Phê duyệt”.")
    add_step_paragraph(doc, "Bước 2", "Click “Xác nhận phê duyệt” trong hộp thoại thông báo.")
    add_paragraph_styled(doc, "Kết quả: Trạng thái phiếu chuyển sang “approved” (Đã duyệt) và bị khóa vĩnh viễn.")
    add_callout(doc, "Tác vụ tự động của hệ thống khi phê duyệt thành công:\n1. Sinh mã QR duy nhất (qr_token): Tạo chuỗi token ngẫu nhiên 64 ký tự gán cho thuộc tính qr_token của từng đại biểu để quét di động bảo mật.\n2. Cấp số thứ tự in thẻ (badge_number): Tự động đánh số thứ tự in thẻ tăng dần (Ví dụ: 001, 002, 003...) chuẩn bị in thẻ vật lý.")
    
    # Summary of Business Rules Table
    add_heading_styled(doc, "6.4 Bảng tổng hợp các ràng buộc nghiệp vụ (Business Rules)", level=2)
    table_br = doc.add_table(rows=8, cols=3)
    headers_br = ["Mã Ràng Buộc", "Nội Dung Ràng Buộc", "Cơ Chế Kiểm Soát (Validation)"]
    col_widths_br = [1.5, 2.3, 3.2]
    data_br = [
        ["BR-REG-01", "Phương thức nhập đại biểu", "Hỗ trợ đồng bộ từ SMILE hoặc tự nhập thủ công nếu SMILE chưa cập nhật."],
        ["BR-REG-02", "Hồ sơ đính kèm bắt buộc", "Mỗi đại biểu bắt buộc phải tải lên đủ 4 file: CCCD Mặt trước, CCCD Mặt sau, Ảnh chân dung, Scan HĐLĐ."],
        ["BR-REG-03", "Kích thước ảnh chân dung", "Validate phía máy chủ: Kích thước ảnh chân dung bắt buộc phải đúng 530x530 pixel."],
        ["BR-REG-04", "Thời gian gia nhập nhân sự", "Tự động chặn đăng ký với nhân viên gia nhập đơn vị từ ngày 01/06/2026 trở đi."],
        ["BR-REG-05", "Giới hạn môn thể thao", "Mỗi đại biểu đăng ký tham gia tối đa N môn thể thao root (mặc định = 3)."],
        ["BR-REG-06", "Phòng ban thi nghiệp vụ", "Đại biểu thi nghiệp vụ bắt buộc phải thuộc mã phòng ban SMILE khớp với cấu hình."],
        ["BR-APPROVE", "Đồng bộ phê duyệt", "Phê duyệt hoặc từ chối thực hiện trên toàn bộ phiếu đăng ký, không phê duyệt riêng lẻ."]
    ]
    format_grid_table(table_br, col_widths_br, headers_br, data_br)
    
    doc.add_page_break()
    
    # ==================== SECTION 7 ====================
    add_heading_styled(doc, "7. Liên hệ hỗ trợ", level=1)
    
    add_heading_styled(doc, "7.1 Bộ phận hỗ trợ kỹ thuật EventRegis", level=2)
    add_bullet_item(doc, "Email: ", "event.support@muongthanh.vn", level=0)
    add_bullet_item(doc, "Hotline: ", "1900 xxxx (nhánh số 2)", level=0)
    add_bullet_item(doc, "Thời gian hỗ trợ: ", "Từ 08:00 đến 17:30 (Thứ 2 đến Thứ 6 hàng tuần)", level=0)
    
    # Save the document
    out_path = r'e:\eventregis\docs\word\Huong_dan_su_dung_EventRegis_theo_mau.docx'
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"User guide generated and saved successfully to {out_path}")

if __name__ == '__main__':
    generate_docx()
