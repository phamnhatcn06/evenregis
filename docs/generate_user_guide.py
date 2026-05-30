import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Paths to actual images in the workspace
IMAGE_PATHS = {
    'login': r'e:\eventregis\docs\images\actual_login.png',
    'dashboard': r'e:\eventregis\docs\images\actual_dashboard.png',
    'view': r'e:\eventregis\docs\images\actual_registration_view.png'
}

# Custom styling helpers
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, space_before=Pt(12), space_after=Pt(6)):
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.space_before = space_before
    heading.paragraph_format.space_after = space_after
    heading.paragraph_format.keep_with_next = True
    
    # Apply Mường Thanh Premium Burgundy color to headings
    run = heading.runs[0]
    run.font.name = 'Arial'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(139, 0, 0)  # Burgundy
        run.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(163, 0, 0)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True
        
    return heading

def add_paragraph_styled(doc, text="", style='Normal', space_after=Pt(6), line_spacing=1.15):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
    return p

def add_bullet_item(doc, bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    run_bold = p.add_run(bold_prefix)
    run_bold.font.name = 'Arial'
    run_bold.font.size = Pt(10.5)
    run_bold.bold = True
    
    run_text = p.add_run(text)
    run_text.font.name = 'Arial'
    run_text.font.size = Pt(10.5)
    return p

def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    
    # Shading background XML
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading)
    
    # Left border XML
    pBrd = OxmlElement('w:pBrd')
    left_border = OxmlElement('w:left')
    left_border.set(qn('w:val'), 'single')
    left_border.set(qn('w:sz'), '24')  # 3pt width
    left_border.set(qn('w:space'), '8')
    left_border.set(qn('w:color'), '8B0000')  # Burgundy border
    pBrd.append(left_border)
    p._p.get_or_add_pPr().append(pBrd)
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.italic = True
    return p

def add_centered_picture(doc, path, caption, width_in=5.8):
    if not path or not os.path.exists(path):
        add_paragraph_styled(doc, f"[Ảnh minh họa không khả dụng: {path}]", style='Normal')
        return
        
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    
    p_img.add_run().add_picture(path, width=Inches(width_in))
    
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(12)
    run_cap = p_cap.add_run(f"Hình minh họa: {caption}")
    run_cap.font.name = 'Arial'
    run_cap.font.size = Pt(9)
    run_cap.italic = True
    run_cap.font.color.rgb = RGBColor(128, 128, 128)

def format_grid_table(table, col_widths, headers, data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header formatting
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], '8B0000')  # Burgundy Header
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Data row formatting
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = 'F9F9F9' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            
    # Set cell widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)

# Main document generation
def generate_docx():
    doc = Document()
    
    # Configure margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(36)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("TÀI LIỆU HƯỚNG DẪN SỬ DỤNG")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(139, 0, 0)
    run_title.bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(48)
    run_sub = p_sub.add_run("QUY TRÌNH ĐĂNG NHẬP & ĐẠI BIỂU ĐĂNG KÝ SỰ KIỆN\n(EVENTREGIS)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.bold = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)
    
    add_paragraph_styled(doc, "Phiên bản: 2.0 (Cập nhật 2026)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_styled(doc, "Tech Stack: Yii1 PHP Framework | Nền tảng Portal SSO").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Section 1
    add_heading_styled(doc, "1. Tổng Quan & Phân Quyền Vai Trò (Actors)", level=1)
    add_paragraph_styled(doc, "Hệ thống EventRegis được xây dựng nhằm số hóa và tự động hóa toàn bộ quy trình tổ chức sự kiện đại hội tập trung với quy mô ~600 đại biểu từ 50-100 đơn vị thành viên. Để đảm bảo tính đồng bộ dữ liệu, kiểm soát gian lận và tối ưu hóa thời gian xử lý, quy trình đăng nhập và đăng ký sự kiện được kiểm soát chặt chẽ thông qua các vai trò (Actors) sau:")
    
    # Table 1: Actor Permissions
    table_actors = doc.add_table(rows=4, cols=3)
    headers_actors = ["Vai trò (Actor)", "Tài khoản sử dụng", "Quyền hạn & Nghiệp vụ chính"]
    col_widths_actors = [1.8, 1.8, 3.4]
    data_actors = [
        ["Đại diện Đơn vị", "Tài khoản đơn vị (unit_accounts)", "Tạo phiếu đăng ký, thiết lập liên quân, đồng bộ/nhập thông tin đại biểu, tải lên hồ sơ pháp lý bắt buộc, đăng ký các hoạt động thể thao/nghiệp vụ và nộp hồ sơ phê duyệt."],
        ["Nhân sự HO (HR)", "Tài khoản nội bộ (users - role=hr)", "Kiểm tra tính hợp lệ của hồ sơ đại biểu (CCCD, ảnh chân dung, HĐLĐ), phê duyệt hoặc từ chối toàn bộ phiếu đăng ký kèm lý do lỗi cụ thể, gán vai trò đại biểu và chỉ định Trưởng đoàn."],
        ["Admin HO", "Tài khoản admin (users - role=admin)", "Toàn quyền cấu hình đợt đăng ký, thiết lập số môn thể thao tối đa cho đại biểu, và quản lý in thẻ đại biểu hàng loạt."]
    ]
    format_grid_table(table_actors, col_widths_actors, headers_actors, data_actors)
    
    # Section 2
    add_heading_styled(doc, "2. Quy Trình Đăng Nhập Hệ Thống (Portal SSO)", level=1)
    add_paragraph_styled(doc, "Hệ thống EventRegis sử dụng giải pháp đăng nhập đơn nhất Single Sign-On (SSO) thông qua hệ thống Portal Mường Thanh nhằm tối ưu bảo mật và đồng nhất thông tin tài khoản nhân viên.")
    
    add_heading_styled(doc, "2.1 Cơ chế Hoạt động của SSO Portal:", level=2)
    add_bullet_item(doc, "Chuyển hướng (Redirect): ", "Khi người dùng truy cập trang chủ http://event.mt:8080/, hệ thống kiểm tra session. Nếu chưa đăng nhập, hệ thống sẽ tự động chuyển hướng sang trang đăng nhập của Portal: https://portal.muongthanh.vn/login?redirect=http://event.mt:8080/")
    add_bullet_item(doc, "Xác thực và Cấp Token: ", "Sau khi người dùng đăng nhập thành công trên Portal, Portal sẽ chuyển hướng ngược lại EventRegis kèm theo một mã token JWT an toàn trong URL: http://event.mt:8080/?sso_token=xxx")
    add_bullet_item(doc, "Xử lý Callback & Lưu Session: ", "Controller SiteController::actionIndex() tiếp nhận sso_token, giải mã chữ ký JWT bằng thuật toán HS256 và mã khóa bí mật JWT_SECRET. Khi token hợp lệ, hệ thống sẽ gọi API /api/sso/me để lấy thông tin chi tiết nhân viên (Mã đơn vị, Phòng ban, Chức vụ) và lưu trữ vào Session của Yii.")
    
    add_heading_styled(doc, "2.2 Quy trình Đăng nhập Thực tế (Khi tài khoản đã đăng nhập sẵn trên Chrome):", level=2)
    add_paragraph_styled(doc, "Trong trường hợp tài khoản của bạn đã được đăng nhập sẵn trên trình duyệt Chrome, hệ thống hỗ trợ đăng nhập nhanh thông qua 2 phương thức cực kỳ tiện lợi:")
    
    add_bullet_item(doc, "Phương thức 1: Đăng nhập từ trang chủ Portal (Portal Dashboard)", "")
    doc.add_paragraph("1. Mở Portal: Người dùng mở trình duyệt Chrome đã lưu session và truy cập trang chủ Portal: https://portal.muongthanh.vn", style='Normal')
    doc.add_paragraph("2. Click chọn Ứng dụng: Tại giao diện màn hình chính của Portal (nơi hiển thị các phân hệ/ứng dụng được phân quyền), tìm và click vào biểu tượng ứng dụng 'Đăng ký Sự kiện' (hoặc 'Event Regis').", style='Normal')
    doc.add_paragraph("3. Tự động chuyển hướng: Hệ thống Portal nhận diện phiên hoạt động, tự động sinh mã Token JWT an toàn và thực hiện chuyển hướng trình duyệt thẳng tới hệ thống EventRegis kèm tham số mã hóa: http://event.mt:8080/?sso_token=<JWT_TOKEN_CỦA_BẠN>", style='Normal')
    doc.add_paragraph("4. Vào thẳng Dashboard: Hệ thống EventRegis tự động xử lý token callback và đưa bạn trực tiếp vào giao diện làm việc chính (http://event.mt:8080/admin/default/index) trong chưa đầy 1 giây mà không cần nhập lại bất kỳ thông tin nào.", style='Normal')
    
    add_bullet_item(doc, "Phương thức 2: Đăng nhập từ trang chủ EventRegis (Silent Login)", "")
    doc.add_paragraph("1. Mở EventRegis: Người dùng nhập trực tiếp địa chỉ trang đăng ký sự kiện trên thanh địa chỉ Chrome: http://event.mt:8080/", style='Normal')
    doc.add_paragraph("2. Click nút Đăng nhập: Hệ thống hiển thị trang đăng nhập (Hình 1), người dùng click vào nút màu tím 'Đăng nhập với Portal'.", style='Normal')
    doc.add_paragraph("3. Nhận diện phiên tự động: Trình duyệt chuyển hướng nhanh sang Portal (https://portal.muongthanh.vn/login?redirect=...). Do Chrome đã lưu sẵn phiên đăng nhập Portal của bạn, hệ thống Portal sẽ nhận diện và xác thực tự động ngay lập tức mà không hiển thị màn hình điền Tên đăng nhập & Mật khẩu.", style='Normal')
    doc.add_paragraph("4. Hoàn tất xác thực: Portal tự động điều hướng ngược trở lại EventRegis kèm theo mã sso_token, hệ thống thiết lập session và cho phép bạn bắt đầu làm việc ngay lập tức.", style='Normal')

    add_heading_styled(doc, "2.3 Quản lý Phiên Đăng nhập (Session Lifecycle):", level=2)
    add_bullet_item(doc, "Thời gian hết hạn (Session Timeout): ", "Phiên làm việc được cấu hình mặc định là 1800 giây (30 phút) không có hoạt động.")
    add_bullet_item(doc, "Thời gian làm mới (Refresh Interval): ", "Mỗi 900 giây (15 phút), hệ thống sẽ tự động làm mới token ngầm để duy trì đăng nhập mà không làm gián đoạn trải nghiệm của người dùng.")
    add_bullet_item(doc, "Đăng xuất (Logout): ", "Khi người dùng click Đăng xuất, SiteController::actionLogout() sẽ hủy toàn bộ session lưu trên server và localStorage trên trình duyệt, sau đó điều hướng người dùng quay trở lại trang đăng nhập.")
    
    add_centered_picture(doc, IMAGE_PATHS['login'], "Giao diện đăng nhập tích hợp Portal SSO chạy tại http://event.mt:8080/", width_in=5.8)

    
    # Section 3
    add_heading_styled(doc, "3. Quy Trình Đăng Ký Sự Kiện (Event Registration)", level=1)
    add_paragraph_styled(doc, "Quy trình đăng ký sự kiện là luồng nghiệp vụ khép kín gồm 5 bước tuần tự bắt buộc dành cho Đại diện Đơn vị:")
    
    add_heading_styled(doc, "Bước 3.1: Khởi tạo Phiếu Đăng Ký", level=2)
    add_paragraph_styled(doc, "Đại diện đơn vị đăng nhập vào hệ thống trong khung thời gian đăng ký được mở (registration_periods). Truy cập menu 'Đăng ký Sự kiện' và click 'Khởi tạo Phiếu Đăng ký'. Hệ thống tự động tạo một phiếu đăng ký với trạng thái mặc định ban đầu là 'draft' (Bản nháp) gắn với đơn vị của tài khoản đó. Mỗi đơn vị chỉ được phép có đúng một phiếu đăng ký hoạt động trong mỗi đợt đăng ký.")
    
    add_heading_styled(doc, "Bước 3.2: Nhập Danh sách Đại Biểu & Tải lên Hồ sơ Bắt buộc", level=2)
    add_paragraph_styled(doc, "Đại diện đơn vị vào trang chi tiết phiếu đăng ký, chọn tab 'Danh sách Đại biểu' -> Click 'Thêm đại biểu' để nhập nhân sự tham gia đại hội.")
    
    add_bullet_item(doc, "A. Nguồn dữ liệu nhân sự (BR-REG-01): ", "Hỗ trợ 2 cơ chế nhập liệu linh hoạt gồm đồng bộ từ SMILE hoặc tự nhập thủ công.")
    doc.add_paragraph("1. Đồng bộ từ Hệ thống SMILE (Khuyên dùng): Người dùng nhập tên hoặc mã nhân viên để tìm kiếm. Hệ thống tự động truy vấn dữ liệu SMILE và điền thông tin: Họ tên, Chức vụ hiện tại, Mã phòng ban. RÀNG BUỘC NGHIỆP VỤ BẮT BUỘC: Hệ thống chỉ cho phép hiển thị và đăng ký những nhân sự có ngày gia nhập đơn vị trước ngày 01/06/2026. Những nhân viên gia nhập từ ngày 01/06/2026 trở đi sẽ bị hệ thống tự động ẩn hoặc chặn không cho đăng ký.", style='Normal')
    doc.add_paragraph("2. Nhập thủ công (Manual CRUD): Tích chọn ô 'Tự nhập thông tin' để điền thủ công các nhân viên chính thức thỏa mãn điều kiện thời gian nhưng chưa cập nhật dữ liệu trên SMILE.", style='Normal')
    
    add_bullet_item(doc, "B. Validate Tài liệu Đính kèm Bắt buộc (BR-REG-02 & BR-REG-03): ", "Mỗi đại biểu bắt buộc phải có đủ 4 file hồ sơ pháp lý sau:")
    doc.add_paragraph("- Ảnh mặt trước CCCD: Định dạng ảnh JPG/PNG, dung lượng tối đa 5MB.", style='List Bullet')
    doc.add_paragraph("- Ảnh mặt sau CCCD: Định dạng ảnh JPG/PNG, dung lượng tối đa 5MB.", style='List Bullet')
    doc.add_paragraph("- Ảnh chân dung in thẻ (Portrait): Bắt buộc phải có kích thước chính xác 530x530 pixel. Validate ở phía server; nếu ảnh sai kích thước, hệ thống sẽ báo lỗi validation và từ chối lưu hồ sơ.", style='List Bullet')
    doc.add_paragraph("- Scan Hợp đồng lao động: File scan PDF hoặc ảnh JPG rõ chữ ký/dấu đỏ thể hiện nhân sự chính thức, tối đa 10MB.", style='List Bullet')
    
    add_callout(doc, "Lưu ý quan trọng: Quy tắc validate ảnh chân dung 530x530px là bắt buộc để ảnh in thẻ đại biểu vật lý đạt tiêu chuẩn, rõ nét, không bị kéo dãn.")
    
    add_heading_styled(doc, "Bước 3.3: Thiết lập Liên quân theo từng Nội dung (Content-level Alliance)", level=2)
    add_paragraph_styled(doc, "Cơ chế Liên quân theo nội dung cho phép các đơn vị nhỏ ghép nhân sự để thành lập đội thi đấu thể thao tập thể hoặc tiết mục văn nghệ chung. Liên quân hoạt động độc lập theo từng nội dung cụ thể, không áp dụng chung cho toàn bộ sự kiện.")
    add_bullet_item(doc, "Gửi yêu cầu liên quân (BR-AL05): ", "Đơn vị gửi yêu cầu chọn đơn vị đối tác, chọn nội dung liên quân (Ví dụ: Bóng đá nam). Trạng thái yêu cầu ở dạng pending (Chờ duyệt).")
    add_bullet_item(doc, "Duyệt yêu cầu liên quân (BR-AL06): ", "Đơn vị đối tác có quyền click Đồng ý (active) hoặc Từ chối kèm lý do.")
    add_bullet_item(doc, "Ràng buộc số lượng (BR-AL02 & BR-AL03): ", "Đơn vị chỉ được phép liên quân tối đa với số lượng đơn vị khác được cấu hình cho mỗi nội dung (max_alliance_orgs).")
    
    add_heading_styled(doc, "Bước 3.4: Đăng ký Hoạt động Chi tiết (Event Registration)", level=2)
    add_paragraph_styled(doc, "Đại diện đơn vị vào phần 'Đăng ký hoạt động' để đăng ký các môn thi đấu:")
    add_bullet_item(doc, "Đăng ký theo số lượng (Quantity-based): ", "Áp dụng cho môn thể thao tập thể. Đơn vị chỉ đăng ký tham gia và số lượng đội thi đấu, chưa cần điền chi tiết danh sách thành viên ở bước này.")
    add_bullet_item(doc, "Đăng ký theo danh sách cụ thể (Detailed-based): ", "Áp dụng cho thi Miss, văn nghệ cá nhân và thi nghiệp vụ. Đơn vị chọn đại biểu cụ thể gán vào nội dung thi đấu.")
    doc.add_paragraph("- Validate môn thể thao tối đa (BR-REG-05): Mỗi đại biểu chỉ được đăng ký tham gia tối đa N môn thể thao root (max_sports_per_attendee, mặc định là 3).", style='List Bullet')
    doc.add_paragraph("- Validate phòng ban thi nghiệp vụ (BR-REG-06): Đại biểu thi nghiệp vụ bắt buộc phải có mã phòng ban thuộc SMILE nằm trong danh mục competition_departments.", style='List Bullet')
    
    add_heading_styled(doc, "Bước 3.5: Nộp Hồ Sơ Đăng Ký", level=2)
    add_paragraph_styled(doc, "Sau khi hoàn thiện danh sách đại biểu, liên quân và nội dung đăng ký, Đại diện đơn vị nhấn nút 'Nộp đăng ký'. Phiếu đăng ký sẽ chuyển trạng thái sang 'submitted'. Hệ thống sẽ tự động khóa toàn bộ quyền chỉnh sửa của đơn vị để đảm bảo tính toàn vẹn dữ liệu trong quá trình phê duyệt.")
    
    add_centered_picture(doc, IMAGE_PATHS['dashboard'], "Giao diện Danh sách phiếu đăng ký đại biểu đơn vị tại http://event.mt:8080/admin/registrations/admin", width_in=5.8)
    
    # Section 4
    add_heading_styled(doc, "4. Quy Trình Kiểm Duyệt Hồ Sơ & Phê Duyệt", level=1)
    add_paragraph_styled(doc, "Nhân sự HO (HR HO) thực hiện quy trình kiểm tra và phê duyệt phiếu đăng ký của các đơn vị trên trang quản trị:")
    
    add_heading_styled(doc, "4.1 Tiếp nhận & Thẩm định Hồ sơ Đại biểu", level=2)
    add_paragraph_styled(doc, "HR HO truy cập trang 'Kiểm duyệt Đăng ký', mở chi tiết phiếu của đơn vị đang có trạng thái 'submitted'. Click xem chi tiết từng đại biểu trong danh sách để kiểm tra tính hợp lệ của ảnh chân dung in thẻ, 2 mặt ảnh CCCD và nội dung tệp scan HĐLĐ để xác minh đại biểu.")
    
    add_heading_styled(doc, "4.2 Quyết định Phê duyệt (Approve / Reject)", level=2)
    add_bullet_item(doc, "Từ chối Phiếu đăng ký (Reject): ", "Nếu phát hiện hồ sơ đại biểu bị lỗi, HR HO click nút 'Từ chối' và bắt buộc phải nhập lý do chi tiết. Phiếu chuyển sang trạng thái 'rejected' và mở khóa quyền chỉnh sửa để đơn vị cập nhật hồ sơ lỗi và nộp lại.")
    add_bullet_item(doc, "Phê duyệt Phiếu đăng ký (Approve): ", "Nếu hồ sơ đạt yêu cầu, HR HO click nút 'Phê duyệt'. Phiếu chuyển sang trạng thái 'approved' và khóa vĩnh viễn.")
    doc.add_paragraph("TÁC VỤ TỰ ĐỘNG CỦA HỆ THỐNG KHI PHÊ DUYỆT:", style='Normal')
    doc.add_paragraph("1. Sinh mã QR duy nhất (qr_token): Tạo một chuỗi token ngẫu nhiên dài 64 ký tự gán cho thuộc tính qr_token của đại biểu dùng để quét tra cứu di động bảo mật.", style='List Bullet')
    doc.add_paragraph("2. Cấp số thứ tự thẻ (badge_number): Tự động sinh số thứ tự in thẻ tăng dần theo sequence (Ví dụ: 001, 002, 003...) chuẩn bị in thẻ vật lý.", style='List Bullet')
    
    add_callout(doc, "Lưu ý quy định: Phê duyệt hoặc từ chối thực hiện trên TOÀN BỘ phiếu đăng ký của đơn vị, không phê duyệt riêng lẻ từng đại biểu để đảm bảo dữ liệu đồng bộ.")
    
    add_centered_picture(doc, IMAGE_PATHS['view'], "Giao diện Chi tiết Phiếu đăng ký của HR HO tại http://event.mt:8080/admin/registrations/view/id/2", width_in=5.8)
    
    # Section 5
    add_heading_styled(doc, "5. Tổng Hợp Các Ràng Buộc Nghiệp Vụ (Business Rules)", level=1)
    
    # Table 2: Business Rules
    table_br = doc.add_table(rows=10, cols=3)
    headers_br = ["Mã Ràng Buộc", "Nội Dung Ràng Buộc", "Cơ Chế Kiểm Soát (Validation)"]
    col_widths_br = [1.5, 2.3, 3.2]
    data_br = [
        ["BR-REG-01", "Nhập danh sách đại biểu", "Cho phép chọn từ SMILE hoặc tự nhập nếu chưa có trên SMILE."],
        ["BR-REG-02", "Hồ sơ đính kèm bắt buộc", "Bắt buộc tải lên đủ 4 file: CCCD Mặt trước, CCCD Mặt sau, Ảnh chân dung, Scan HĐLĐ."],
        ["BR-REG-03", "Kích thước ảnh chân dung", "Validate phía máy chủ: Kích thước ảnh chân dung bắt buộc phải đúng 530x530 pixel."],
        ["BR-REG-04", "Thời gian gia nhập nhân viên", "Tự động ẩn hoặc chặn đăng ký với nhân viên gia nhập đơn vị từ ngày 01/06/2026 trở đi."],
        ["BR-REG-05", "Giới hạn môn thể thao", "Mỗi đại biểu đăng ký tối đa N môn thể thao root (Mặc định max_sports_per_attendee = 3)."],
        ["BR-REG-06", "Phòng ban thi nghiệp vụ", "Đại biểu thi nghiệp vụ phải có mã phòng ban SMILE khớp với cấu hình competition_departments."],
        ["BR-AL01", "Điều kiện liên quân nội dung", "Chỉ áp dụng liên quân cho các môn thể thao tập thể có cấu hình cho phép liên quân."],
        ["BR-AL03", "Giới hạn số đơn vị liên quân", "Số lượng đơn vị ghép đội liên quân không được vượt quá max_alliance_orgs cấu hình cho nội dung đó."],
        ["BR-APPROVE", "Đồng bộ trạng thái phê duyệt", "Phê duyệt hoặc từ chối thực hiện trên toàn bộ phiếu đăng ký của đơn vị, không phê duyệt riêng lẻ."]
    ]
    format_grid_table(table_br, col_widths_br, headers_br, data_br)
    
    # Save the document
    out_path = r'e:\eventregis\docs\word\Huong_dan_su_dung_EventRegis.docx'
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"User guide generated and saved successfully to {out_path}")

if __name__ == '__main__':
    generate_docx()
