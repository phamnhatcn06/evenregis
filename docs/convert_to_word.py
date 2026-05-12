import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_code_block(doc, code, language=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)

    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    p._p.get_or_add_pPr().append(shading)

def process_table(doc, lines):
    rows = []
    for line in lines:
        if line.strip().startswith('|') and not line.strip().startswith('|---'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(cells)

    if not rows:
        return

    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < cols:
                table.rows[i].cells[j].text = cell_text
                if i == 0:
                    table.rows[i].cells[j].paragraphs[0].runs[0].bold = True

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    styles = doc.styles
    style = styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            process_table(doc, table_lines)
            table_lines = []
            in_table = False

        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(text.strip(), style='List Number')
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        elif line.strip():
            text = line.strip()
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            doc.add_paragraph(text)
        else:
            pass

        i += 1

    if in_table and table_lines:
        process_table(doc, table_lines)

    doc.save(docx_path)
    print(f'Saved to {docx_path}')

if __name__ == '__main__':
    md_to_docx(
        'e:/eventregis/docs/system-design.md',
        'e:/eventregis/docs/word/system-design.docx'
    )
